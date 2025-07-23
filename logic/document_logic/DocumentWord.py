import json
import os
from docx import Document


from config.settings import Config

#Call to the general configuration container dictionary
if os.path.exists("config/config_program.json"):
    config = Config()
    config.save_config_json()
with open("config/config_program.json", "r", encoding="utf-8") as file:
    config_ = json.load(file)


class WordDocument:
    """
    Class to generate a .docx file from Wikipedia scrapped data.
    """
    def _init_(self, wp_page):
        """
        Initialize the Document class with an instance of the Wikipedia class.
        
        """
        self.wp_page = wp_page
        self.name = wp_page.name
        self.title = wp_page.principal_title.text
        self.json_file = os.path.abspath(
            os.path.join("data_json", f"{self.name}_data.json"))


        self.docx_file = f"{self.name}.docx"

    def read_json_data(self):
        """
        Read data from the corresponding JSON file.
        """
        
        if not os.path.exists(self.json_file):
            raise FileNotFoundError(
                f"JSON file '{self.json_file}' not found.")
        
        with open(self.json_file, "r", encoding="utf-8") as file:
            return json.load(file)

    def add_paragraph(self, doc, text):

        if text and text.strip():
            doc.add_paragraph(text.strip())

    def add_table(self, doc, table_data):
        """
        Add a table to the document from the JSON table data.
        Handles both list-of-lists and list-of-dictionaries.
        """
        if not table_data:
            return


        if isinstance(table_data[0], dict):
            columns = []
            for row in table_data:
                for key in row.keys():
                    if key not in columns and key is not None:
                                columns.append(key)
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = "Table Grid"

            hdr_cells = table.rows[0].cells
            for idx, col_name in enumerate(columns):
                hdr_cells[idx].text = col_name

            for row_data in table_data:
                row_cells = table.add_row().cells
                for idx, col_name in enumerate(columns):
                    value = row_data.get(col_name, "")
                    row_cells[idx].text = str(
                        value) if value is not None else ""

        elif isinstance(table_data[0], list):
            table = doc.add_table(rows=len(table_data), 
                                cols=len(table_data[0]))
            for row_idx, row in enumerate(table_data):
                for col_idx, cell in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(cell)

        else:
            print("Unsupported table format")


    def process_data(self, doc, data):
        """
        Process the data and adds it to the document.
        """
        for key, value in data.items():
            if isinstance(value,list):
                for item in value:
                    if isinstance(item,dict):
                        for sub_key, sub_value in item.items():
                            if "text_p" in sub_key:
                                self.add_paragraph(
                                doc, sub_value.get(sub_key,""))
                            elif "text_h2" in sub_key:
                                doc.add_heading(
                                sub_value.get(sub_key,""), level=2)
                            elif "list_ul" in sub_key:
                                if isinstance(
                                    sub_value,dict) and 'li' in sub_value:
                                    for li_item in sub_value['li']:
                                        doc.add_paragraph(
                                        li_item,style='List Bullet')
                            elif "table_table" in sub_key:
                                self.add_table(doc, sub_value)
            else:
                self.add_paragraph(doc, value)

    def generate(self):
        """
        Create a .docx file from the JSON data.
        """
        doc = Document()
        doc.add_heading(self.title,0)
        data = self.read_json_data()
        self.process_data(doc,data)

        try:
            doc.save(self.docx_file)
            print(f"Document '{self.docx_file}' created successfully")
        except PermissionError:
            print(
        f"Cannot overwrite '{self.docx_file}'. Is it open or locked?")